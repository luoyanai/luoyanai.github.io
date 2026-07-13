from pathlib import Path
import subprocess
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "tmp_vhs_docx"
OUT = ROOT / "output"
WORK.mkdir(exist_ok=True)
OUT.mkdir(exist_ok=True)

md = r'''# 3.3.3　变高度稳定化

::: {custom-style="BlueModified"}
前述多步可捕获性分析主要通过落足位置和DCM偏置调节扩大机器人在水平面内的可恢复稳定域。然而，当机器人由平地进入坡面、连续起伏地形或动态翻转支撑面时，支撑面倾角和落足高度会发生变化，仅依靠水平落足调节难以同时保证摆动足离地间隙、落足安全性和腿部运动学可行性。为此，本文引入变高度稳定化（variable height stabilization，VHS）方法，根据支撑条件变化调节质心高度基准和摆动步高基准，为DCM多步规划补充垂向步态参数调节能力。
:::

::: {custom-style="BlueModified"}
VHS输出质心高度基准 $h_c^{r}$ 和摆动步高基准 $h_s^{r}$。其中，质心高度影响腿部屈伸构型、系统动态响应和落足过程中的稳定裕度；摆动步高决定摆动足竖直轨迹幅值，直接影响坡面通过过程中的离地间隙和落足安全性。高度设置过小可能造成腿部活动空间不足或摆动足擦碰，高度设置过大则会增加身体姿态波动和关节执行负担。因此，需要针对不同坡度条件确定相应的质心高度和摆动步高基准。
:::

::: {custom-style="BlueModified"}
本文采用“离线构建、在线调用”的方式建立VHS曲线。离线阶段，在Gazebo仿真环境中构建坡度角为 $-15^\circ$～$20^\circ$ 的固定坡面，并以 $5^\circ$ 为间隔设置代表性坡度工况，将地形坡度角记为 $\theta$。对于每一个坡度工况，分别改变质心高度 $h_c$ 和摆动步高 $h_s$，采用二分搜索逐步缩小稳定参数范围，并针对各参数组合开展10～20次坡面行走测试，由此确定该坡度下稳定行走的质心高度范围、摆动步高范围及综合表现较优的参数点。
:::

::: {custom-style="BlueModified"}
为评价不同高度组合下的坡面行走表现，定义最大稳定通过距离 $X_\theta(h_c,h_s)$ 和最长稳定行走时间 $T_{\mathrm{stable}}(h_c,h_s)$，并以二者的乘积作为同一坡度工况下的行走效率指标：
:::

::: {custom-style="BlueModified"}
$$P=X_\theta(h_c,h_s)\,T_{\mathrm{stable}}(h_c,h_s)$$
:::

::: {custom-style="BlueEquationNumber"}
（3-30）
:::

::: {custom-style="BlueModified"}
式中，$X_\theta(h_c,h_s)$ 表示机器人在坡度角 $\theta$ 下采用质心高度 $h_c$ 和摆动步高 $h_s$ 时能够达到的最大稳定通过距离；$T_{\mathrm{stable}}(h_c,h_s)$ 表示相同参数条件下能够维持稳定行走的最长时间。对于每一个坡度角，选取行走效率 $P$ 较高且能够稳定完成坡面通过的质心高度和摆动步高作为该坡度下的高度基准点，同时记录其稳定可行区间。由此获得地形坡度角与质心高度基准、摆动步高基准之间的离散对应数据。
:::

::: {custom-style="BlueModified"}
为便于在线计算并减少离散查表引起的参数跳变，采用MATLAB分别尝试线性函数、多项式函数、样条函数等模型对离散高度基准点进行拟合。综合比较拟合误差和曲线连续性后，采用有理函数描述摆动步高基准随坡度的变化关系，采用傅里叶函数描述质心高度基准随坡度的变化关系，得到图3-6所示的VHS曲线：
:::

$$h_s^{r}(\theta)=\frac{\alpha_1\theta^2+\alpha_2\theta+\alpha_3}{\theta^2+\beta_1\theta+\beta_2}$$

::: {custom-style="EquationNumber"}
（3-31）
:::

$$h_c^{r}(\theta)=A_0+A_1\cos(K_s\theta)+B_1\sin(K_s\theta)$$

::: {custom-style="EquationNumber"}
（3-32）
:::

::: {custom-style="BlueModified"}
式中，$h_s^{r}(\theta)$ 为坡度角 $\theta$ 对应的摆动步高基准，$h_c^{r}(\theta)$ 为对应的质心高度基准；$\alpha_1$、$\alpha_2$、$\alpha_3$、$\beta_1$、$\beta_2$、$A_0$、$A_1$、$B_1$ 和 $K_s$ 为曲线拟合参数，具体取值见表3-1。式（3-31）和式（3-32）中的 $\theta$ 按离线拟合时采用的角度数值输入，在线坡度估计结果应统一到相同的零位、正负方向和角度单位。表中高度参数以厘米给出，进入控制器前统一换算为机器人内部使用的长度单位。
:::

**表3-1　G1机器人VHS曲线拟合参数**

**Table 3-1　Fitting parameters of the VHS curves for the G1 robot**

| 基准高度（cm） | 拟合参数 | 参数值 | 决定系数 $R^2$ |
|:---:|:---:|:---:|:---:|
| $h_c^{r}(\theta)$ | $A_0$ | 88.202 | 0.8953 |
| $h_c^{r}(\theta)$ | $A_1$ | -6.299 | 0.8953 |
| $h_c^{r}(\theta)$ | $B_1$ | 6.521 | 0.8953 |
| $h_c^{r}(\theta)$ | $K_s$ | 1.495 | 0.8953 |
| $h_s^{r}(\theta)$ | $\alpha_1$ | 6.398 | 0.9996 |
| $h_s^{r}(\theta)$ | $\alpha_2$ | 0.742 | 0.9996 |
| $h_s^{r}(\theta)$ | $\alpha_3$ | 0.149 | 0.9996 |
| $h_s^{r}(\theta)$ | $\beta_1$ | 0.617 | 0.9996 |
| $h_s^{r}(\theta)$ | $\beta_2$ | 0.337 | 0.9996 |

::: {custom-style="BlueModified"}
在线运行时，机器人无法预先获得时变支撑面的准确坡度，因此VHS利用踝关节状态估计当前支撑条件。支撑面坡度改变或平台快速翻转时，踝关节俯仰角 $\theta_{\mathrm{ankle}}$ 随足底姿态发生变化，其角速度 $v_\theta$ 则在突变阶段明显增大。首先将机器人关节坐标系下的踝关节俯仰角结合当前机器人姿态转换至世界坐标系，得到与离线坡度标定一致的当前坡度估计：
:::

::: {custom-style="BlueModified"}
$$\theta=\operatorname{RobotToWorldFrame}(\theta_{\mathrm{ankle}})$$
:::

::: {custom-style="BlueModified"}
该过程并非将踝关节角与地形坡度简单设为相等，而是将足底俯仰姿态统一到世界坐标系下，从而获得可用于调用离线VHS曲线的坡度量。当踝关节角速度幅值超过设定阈值 $v_\theta^{\mathrm{th}}$ 时，判定机器人正在经历明显的支撑条件变化，并根据当前坡度估计更新高度基准：
:::

$$h_c\leftarrow h_c^{r}(\theta),\qquad h_s\leftarrow h_s^{r}(\theta),\qquad |v_\theta|\geq v_\theta^{\mathrm{th}}$$

::: {custom-style="EquationNumber"}
（3-33）
:::

::: {custom-style="BlueModified"}
本文将扰动触发阈值设为 $10\ \mathrm{rad/s}$。若踝关节角速度未达到阈值，则维持上一周期的质心高度和摆动步高基准，避免正常周期性行走过程中频繁重构轨迹。为抑制测量噪声和瞬时尖峰引起的高度指令突变，对单周期高度变化量进一步施加限幅：
:::

$$h_c\leftarrow\operatorname{sat}\!\left(h_c^{r}(\theta),\,h_c^{\mathrm{pre}}\pm\Delta h_c^{\max}\right),\qquad h_s\leftarrow\operatorname{sat}\!\left(h_s^{r}(\theta),\,h_s^{\mathrm{pre}}\pm\Delta h_s^{\max}\right)$$

::: {custom-style="EquationNumber"}
（3-34）
:::

::: {custom-style="BlueModified"}
式中，$h_c^{\mathrm{pre}}$ 和 $h_s^{\mathrm{pre}}$ 为上一控制周期的高度基准，$\Delta h_c^{\max}$ 和 $\Delta h_s^{\max}$ 为单周期允许的最大变化量，$\operatorname{sat}(\cdot)$ 为限幅算子。更新后的质心高度基准用于调整倒立摆等效高度及DCM动态时间尺度，摆动步高基准用于确定摆动足竖直轨迹幅值。在此基础上，多步可捕获规划进一步优化落足位置、DCM偏置和相位参数，从而形成垂向高度调节与水平稳定恢复相结合的动态稳定控制过程。
:::
'''

md_path = WORK / "vhs.md"
md_path.write_text(md, encoding="utf-8")

# Build a reference Word document with thesis-like styles.
ref = Document()
sec = ref.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.6)

styles = ref.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.first_line_indent = Cm(0.84)

h1 = styles["Title"]
h1.font.name = "Times New Roman"
h1.font.size = Pt(16)
h1.font.bold = True
h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(12)

blue = styles.add_style("BlueModified", 1)
blue.base_style = normal
blue.font.name = "Times New Roman"
blue.font.size = Pt(12)
blue.font.color.rgb = RGBColor(0, 0, 255)
blue._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
blue.paragraph_format.line_spacing = 1.5
blue.paragraph_format.space_after = Pt(0)
blue.paragraph_format.first_line_indent = Cm(0.84)

blue_num = styles.add_style("BlueEquationNumber", 1)
blue_num.base_style = normal
blue_num.font.color.rgb = RGBColor(0, 0, 255)
blue_num._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
blue_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
blue_num.paragraph_format.first_line_indent = Cm(0)

eq_num = styles.add_style("EquationNumber", 1)
eq_num.base_style = normal
eq_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
eq_num.paragraph_format.first_line_indent = Cm(0)

ref_path = WORK / "reference.docx"
ref.save(ref_path)

out_path = OUT / "VHS_3.3.3_editable.docx"
subprocess.run([
    "pandoc", str(md_path), "-o", str(out_path),
    "--reference-doc", str(ref_path),
    "--from", "markdown+tex_math_dollars+fenced_divs+bracketed_spans",
    "--standalone"
], check=True)

# Post-process layout, fonts, equation-number alignment, and the parameter table.
doc = Document(out_path)
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.6)

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt == "3.3.3　变高度稳定化":
        p.style = doc.styles["Title"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
    if txt in {"（3-30）"}:
        p.style = doc.styles["BlueEquationNumber"]
    elif txt in {"（3-31）", "（3-32）", "（3-33）", "（3-34）"}:
        p.style = doc.styles["EquationNumber"]
    if txt.startswith("表3-1") or txt.startswith("Table 3-1"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if p.style.name == "Title":
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.bold = True
            run.font.size = Pt(16)

# Format and merge the VHS parameter table.
for table in doc.tables:
    if table.rows and table.cell(0, 0).text.strip() == "基准高度（cm）":
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = [Cm(4.0), Cm(3.0), Cm(3.0), Cm(3.0)]
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = widths[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10.5)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        # Merge the function and R^2 cells to match the original table structure.
        table.cell(1, 0).merge(table.cell(4, 0))
        table.cell(1, 3).merge(table.cell(4, 3))
        table.cell(5, 0).merge(table.cell(9, 0))
        table.cell(5, 3).merge(table.cell(9, 3))
        # Set borders.
        tblPr = table._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = qn(f"w:{edge}")
            el = borders.find(tag)
            if el is None:
                el = OxmlElement(f"w:{edge}")
                borders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "000000")

# Add document properties.
doc.core_properties.title = "3.3.3 变高度稳定化修改稿"
doc.core_properties.subject = "按照小论文方法重写的VHS章节，公式为可编辑Office Math"
doc.core_properties.author = ""
doc.save(out_path)
print(out_path)
